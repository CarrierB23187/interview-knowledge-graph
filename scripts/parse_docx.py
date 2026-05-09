"""Parse Word (.docx) files into structured chapters with text and images."""

import base64
import json
import os
from pathlib import Path
from io import BytesIO

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def extract_images_from_docx(doc: Document) -> list[dict]:
    """Extract images embedded in a Word document."""
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_blob = rel.target_part.blob
            ext = rel.target_ext
            b64 = base64.b64encode(image_blob).decode("utf-8")
            mime_map = {
                "png": "image/png",
                "jpg": "image/jpeg",
                "jpeg": "image/jpeg",
                "gif": "image/gif",
                "bmp": "image/bmp",
                "webp": "image/webp",
                "svg": "image/svg+xml",
            }
            images.append({
                "ext": ext,
                "base64": b64,
                "mime": mime_map.get(ext, f"image/{ext}"),
            })
    return images


def is_heading(paragraph) -> bool:
    """Check if a paragraph is a heading."""
    return paragraph.style.name.startswith("Heading") if paragraph.style else False


def get_heading_level(paragraph) -> int:
    """Get the heading level (1-6) of a paragraph."""
    if not paragraph.style or not paragraph.style.name.startswith("Heading"):
        return 0
    try:
        return int(paragraph.style.name.replace("Heading", "").strip())
    except ValueError:
        # Handle "Heading 1" → "Heading1" variants
        name = paragraph.style.name
        digits = "".join(c for c in name if c.isdigit())
        return int(digits) if digits else 1


def group_by_headings(doc: Document) -> list[dict]:
    """Group document content by heading hierarchy."""
    chapters = []
    current_chapter = {"title": "前言", "text": [], "level": 0}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if is_heading(para):
            # Save previous chapter
            if current_chapter["text"]:
                current_chapter["text"] = "\n".join(current_chapter["text"])
                chapters.append(current_chapter)

            level = get_heading_level(para)
            current_chapter = {"title": text, "text": [], "level": level}
        else:
            current_chapter["text"].append(text)

    # Save last chapter
    if current_chapter["text"]:
        current_chapter["text"] = "\n".join(current_chapter["text"])
        chapters.append(current_chapter)

    return chapters


def parse_docx(filepath: str) -> dict:
    """Parse a Word document and return structured content.

    Returns:
        {
            "fileName": "xxx.docx",
            "fullText": "...",
            "chapters": [
                {"title": "...", "text": "...", "level": N}
            ]
        }
    """
    doc = Document(filepath)
    filename = Path(filepath).name

    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    chapters = group_by_headings(doc)

    # Remove the 'level' key from output for cleaner API
    images = extract_images_from_docx(doc)
    for ch in chapters:
        ch.pop("level", None)
        ch["images"] = images if chapters.index(ch) == 0 else []
        # Only attach images to the first chapter to avoid duplication

    return {
        "fileName": filename,
        "fullText": full_text,
        "chapters": chapters,
    }


def parse_docx_in_folder(folder: str) -> list[dict]:
    """Parse all .docx files in a folder."""
    results = []
    for f in sorted(Path(folder).glob("*.docx")):
        print(f"  Parsing DOCX: {f.name}")
        result = parse_docx(str(f))
        results.append(result)
    return results


if __name__ == "__main__":
    import sys

    folder = sys.argv[1] if len(sys.argv) > 1 else "../八股"
    output = sys.argv[2] if len(sys.argv) > 2 else "../output/parsed_docs.json"

    results = parse_docx_in_folder(folder)
    os.makedirs(os.path.dirname(output), exist_ok=True)
    with open(output, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    print(f"\nParsed {len(results)} DOCX files → {output}")
